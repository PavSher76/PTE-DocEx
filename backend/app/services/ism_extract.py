"""Извлечение текста из документов ИСМ (PDF, DOC/DOCX, XLS/XLSX)."""

from __future__ import annotations

import logging
import shutil
import subprocess
from pathlib import Path
from tempfile import TemporaryDirectory

from openpyxl import load_workbook

from app.config import Settings
from app.services.pdf_text import extract_pdf_text

logger = logging.getLogger(__name__)

ISM_ALLOWED_SUFFIXES = frozenset({".pdf", ".doc", ".docx", ".xls", ".xlsx"})


def is_ism_allowed_filename(name: str) -> bool:
    return Path(name).suffix.lower() in ISM_ALLOWED_SUFFIXES


def extract_ism_text(path: Path, settings: Settings) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path, settings).text.strip()
    if suffix == ".xlsx":
        return _extract_xlsx(path)
    if suffix == ".xls":
        return _extract_via_libreoffice_text(path)
    if suffix in {".doc", ".docx"}:
        if suffix == ".docx":
            text = _try_extract_docx(path)
            if text.strip():
                return text
        return _extract_via_libreoffice_text(path)
    raise ValueError(f"Неподдерживаемый формат: {suffix}")


def _try_extract_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        return ""
    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    wb = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"## {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts).strip()


def _extract_via_libreoffice_text(path: Path) -> str:
    if not shutil.which("libreoffice"):
        raise RuntimeError(
            "Для .doc/.xls нужен LibreOffice в контейнере backend "
            "(SKIP_APT_PACKAGES=0 при сборке образа)."
        )
    with TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        command = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "txt:Text",
            "--outdir",
            str(tmp_dir),
            str(path),
        ]
        completed = subprocess.run(command, capture_output=True, text=True, timeout=120, check=False)
        if completed.returncode != 0:
            raise RuntimeError(f"LibreOffice: {completed.stderr or completed.stdout}")
        converted = tmp_dir / f"{path.stem}.txt"
        if not converted.exists():
            candidates = list(tmp_dir.glob("*.txt"))
            if not candidates:
                raise RuntimeError("LibreOffice не создал текстовый файл.")
            converted = candidates[0]
        return converted.read_text(encoding="utf-8", errors="replace").strip()
